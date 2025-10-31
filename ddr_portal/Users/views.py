# your_app/views.py

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.db import connection, transaction
from django.db.models import Q
from django.http import HttpResponse, HttpResponseForbidden, FileResponse, Http404, JsonResponse
from django.urls import reverse
from django.utils import timezone
from django.utils.timezone import now
from django.views.decorators.http import require_POST

# Third-party libraries
from openpyxl import Workbook
from openpyxl.styles import Font
from docx import Document as DocxDocument
import pandas as pd
import csv
import openpyxl, base64
import io
import json
from io import BytesIO
import hashlib
import mimetypes

# Local imports
from .forms import CustomUserCreationForm, CustomLoginForm
from .models import (
    Folder, Document, Upload, Role, FolderUserRole,
    UserRole, ActivityLog
)
from .utils import get_user_role_id



# -----------------------------
# Auth-related views
# -----------------------------

def signup_view(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('role_redirect')
    else:
        form = CustomUserCreationForm()
    return render(request, 'signup.html', {'form': form})


def login_view(request):
    if request.method == 'POST':
        form = CustomLoginForm(data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('/role_redirect')
    else:
        form = CustomLoginForm()
    return render(request, 'login.html', {'form': form})


def logout_view(request):
    logout(request)
    return redirect('login')


# -----------------------------
# Role-based dashboards
# -----------------------------

def get_user_role(user_id):
    with connection.cursor() as cursor:
        cursor.execute("SELECT role_id FROM USER_ROLES WHERE user_id = %s", [user_id])
        row = cursor.fetchone()
    if row:
        return row[0]
    return None

@login_required
def admin_view(request):
    role = get_user_role(request.user.id)
    if role != 1:
        return HttpResponseForbidden("You do not have permission to access this page.")
    return render(request, 'admin.html')

@login_required
def head_view(request):
    role = get_user_role(request.user.id)
    if role != 2:
        return HttpResponseForbidden("You do not have permission to access this page.")
    return render(request, 'head.html')

@login_required
def faculty_view(request):
    role = get_user_role(request.user.id)
    if role != 3:
        return HttpResponseForbidden("You do not have permission to access this page.")
    return render(request, 'faculty.html')

@login_required
def role_redirect(request):
    user_id = request.user.id
    with connection.cursor() as cursor:
        cursor.execute("SELECT role_id FROM USER_ROLES WHERE user_id = %s", [user_id])
        row = cursor.fetchone()

    if row:
        role_id = row[0]
        if role_id == 1:
            return redirect('admin')
        elif role_id == 2:
            return redirect('head')
        elif role_id == 3:
            return redirect('faculty')

    return redirect('login')

# -----------------------------
# Folder and document views
# -----------------------------

@login_required
def folder_list(request):
    # Only root folders (parent is NULL)
    folders = Folder.objects.filter(is_active=True, parent__isnull=True)
    return render(request, 'folder_list.html', {'folders': folders})

@login_required
def folder_documents(request, folder_id):
    folder = get_object_or_404(Folder, id=folder_id, is_active=True)

    # Documents inside this folder
    documents = Document.objects.filter(
        folder=folder,
        is_deleted=False,
        folder__is_active=True
    )

    # Subfolders inside this folder
    subfolders = Folder.objects.filter(
        parent=folder,
        is_active=True
    )

    return render(request, 'documents.html', {
        'folder': folder,
        'documents': documents,
        'subfolders': subfolders
    })

# -----------------------------
# Excel file download view
# -----------------------------

@login_required
def download_excel(request, doc_id):
    document = get_object_or_404(
        Document,
        id=doc_id,
        is_deleted=False,
        folder__is_active=True
    )

    raw_data = document.dynamic_data

    try:
        # Clean newline characters
        if isinstance(raw_data, str):
            raw_data = raw_data.replace('\n', '').replace('\r', '').strip()
        else:
            return HttpResponse("dynamic_data is not a string", status=400)

        # Parse JSON
        dynamic_data = json.loads(raw_data)

        # Get main key and columns
        first_key = next(iter(dynamic_data))
        columns = dynamic_data[first_key]["columns"]

        # Create Excel workbook
        wb = Workbook()
        ws = wb.active
        ws.title = document.title
        ws.sheet_properties.tabColor = "1072BA"
        ws.freeze_panes = "A2"

        # Write column headers
        for index, col in enumerate(columns, start=1):
            name = col.get("name", "Unnamed")
            cell = ws.cell(row=1, column=index, value=name)
            cell.font = Font(bold=True)
            ws.column_dimensions[cell.column_letter].width = max(len(name) + 2, 20)

        # Prepare HTTP response
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        filename = f"{document.title.replace(' ', '_')}.xlsx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        wb.save(response)
        return response

    except json.JSONDecodeError as je:
        print("JSON Decode Error:", je)
        return HttpResponse(f"JSON Decode Error: {je}", status=400)

    except KeyError as ke:
        print("Missing expected key:", ke)
        return HttpResponse(f"Missing expected key: {ke}", status=400)

    except Exception as e:
        print("Unexpected Error:", e)
        return HttpResponse(f"Unexpected Error: {e}", status=400)

@login_required
def profile_view(request):
    user = request.user

    try:
        user_role = UserRole.objects.select_related('role').get(user=user)
        role_name = user_role.role.role_name
    except UserRole.DoesNotExist:
        role_name = 'Not Assigned'

    return render(request, 'profile.html', {
        'name': user.get_full_name() or user.username,
        'email': user.email,
        'phone': 'Not Available',  # or customize if needed
        'role': role_name,
    })

def home_view(request):
    return render(request, 'home.html')

@login_required
def see_template(request, doc_id):
    document = get_object_or_404(
        Document,
        id=doc_id,
        is_deleted=False,
        folder__is_active=True
    )
    raw_data = document.dynamic_data

    try:
        if isinstance(raw_data, str):
            raw_data = raw_data.replace('\n', '').replace('\r', '').strip()
        else:
            return HttpResponse("dynamic_data is not a string", status=400)

        dynamic_data = json.loads(raw_data)
        first_key = next(iter(dynamic_data))
        columns = dynamic_data[first_key]["columns"]

        return render(request, 'template_preview.html', {
            'document': document,
            'columns': columns
        })

    except json.JSONDecodeError as je:
        return HttpResponse(f"JSON Decode Error: {je}", status=400)
    except KeyError as ke:
        return HttpResponse(f"Missing key: {ke}", status=400)
    except Exception as e:
        return HttpResponse(f"Unexpected error: {e}", status=400)

    except json.JSONDecodeError as je:
        return HttpResponse(f"JSON Decode Error: {je}", status=400)
    except KeyError as ke:
        return HttpResponse(f"Missing key: {ke}", status=400)
    except Exception as e:
        return HttpResponse(f"Unexpected error: {e}", status=400)

# -----------------------------
# Uploaded Section
#-----------------------------


@login_required
def upload_folder_list(request):
    user = request.user

    # Check if user is admin
    is_admin = FolderUserRole.objects.filter(
        user=user,
        role__role_name__iexact="admin"
    ).exists()

    if is_admin:
        # Admin sees all root folders
        folders = Folder.objects.filter(parent__isnull=True, is_active=True)
    else:
        # Non-admin sees only root folders they have access to
        folders = Folder.objects.filter(
            id__in=FolderUserRole.objects.filter(user=user).values_list('folder_id', flat=True),
            parent__isnull=True,
            is_active=True
        ).distinct()

    return render(request, 'upload_file_list.html', {'folders': folders})

@login_required
def upload_document_list(request, folder_id):
    user = request.user

    # ✅ Step 1: Check if the user is admin
    is_admin = FolderUserRole.objects.filter(
        user=user,
        role__role_name__iexact="admin"
    ).exists()

    # ✅ Step 2: Check folder access (non-admins must be assigned to the folder)
    if not is_admin:
        has_access = FolderUserRole.objects.filter(user=user, folder_id=folder_id).exists()
        if not has_access:
            return HttpResponse("Unauthorized", status=403)

    # ✅ Step 3: Get the folder (must be active)
    folder = get_object_or_404(Folder, id=folder_id, is_active=True)

    # ✅ Step 4: Show only assigned documents
    # Admins see all documents, others see only those assigned to them
    if is_admin:
        documents = Document.objects.filter(folder=folder, is_deleted=False)
    else:
        assigned_file_ids = FolderUserRole.objects.filter(
            user=user,
            folder=folder
        ).values_list('file_id', flat=True)
        documents = Document.objects.filter(
            id__in=assigned_file_ids,
            is_deleted=False
        )

    # ✅ Step 5: Subfolders logic (same as before)
    if is_admin:
        subfolders = Folder.objects.filter(parent=folder, is_active=True)
    else:
        subfolders = Folder.objects.filter(
            parent=folder,
            is_active=True,
            id__in=FolderUserRole.objects.filter(user=user).values_list('folder_id', flat=True)
        ).distinct()

    # ✅ Step 6: Render page
    return render(request, 'upload_document_list.html', {
        'folder': folder,
        'documents': documents,
        'subfolders': subfolders
    })


@require_POST

@login_required
def ajax_upload_file(request):
    try:
        uploaded_file = request.FILES.get('file')
        folder_id = request.POST.get('folder_id')
        document_id = request.POST.get('document_id')

        if not uploaded_file or not folder_id or not document_id:
            return JsonResponse({'success': False, 'error': 'Missing data'})

        folder = get_object_or_404(Folder, id=folder_id, is_active=True)
        document = get_object_or_404(Document, id=document_id, is_deleted=False)
        file_blob = uploaded_file.read()

        upload_instance = Upload.objects.create(
            document=document,
            folder=folder,
            uploaded_by=request.user,
            file_name=uploaded_file.name,
            file_blob=file_blob,
            file_size=uploaded_file.size,
            mime_type=uploaded_file.content_type,
            sha256_hash=hashlib.sha256(file_blob).hexdigest()
        )
        ActivityLog.objects.create(
            user=request.user,
            upload=upload_instance,
            file_name=upload_instance.file_name,
            action="UPLOAD"
        )


        return JsonResponse({'success': True})

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

def get_visible_documents(user, folder):
    # Get roles for the user in this folder
    user_roles = FolderUserRole.objects.filter(user=user, folder=folder)

    # If the user is Head, return all documents
    if user_roles.filter(role__name='Head').exists():
        return Document.objects.filter(folder=folder, is_active=True)

    # If Faculty, return only assigned documents
    faculty_roles = user_roles.filter(role__name='Faculty', file__isnull=False, is_active=True, is_deleted=False)
    return Document.objects.filter(id__in=faculty_roles.values_list('file_id', flat=True), is_active=True, is_deleted=False)

@login_required
def assign_folder_role(request):
    from django.utils.timezone import now

    admin_role_id = Role.objects.get(role_name__iexact="admin").role_id

    def assign_recursive(user_id, folder_id, role_id):
        """
        Recursively assign role to folder, all its subfolders, and documents.
        """
        # Assign role to folder itself
        FolderUserRole.objects.get_or_create(
            user_id=user_id,
            folder_id=folder_id,
            role_id=role_id,
            defaults={"assigned_at": now()}
        )

        # Assign role to all documents inside this folder
        documents = Document.objects.filter(folder_id=folder_id, is_deleted=False)
        for doc in documents:
            FolderUserRole.objects.get_or_create(
                user_id=user_id,
                file_id=doc.id,
                role_id=role_id,
                folder_id=folder_id,
                defaults={"assigned_at": now()}
            )

        # Recurse into subfolders
        subfolders = Folder.objects.filter(parent_id=folder_id, is_active=True)
        for sub in subfolders:
            assign_recursive(user_id, sub.id, role_id)

    if request.method == "POST":
        user_id = request.POST.get("user_id")
        folder_id = request.POST.get("folder_id")
        role_id = int(request.POST.get("role_id"))
        file_id = request.POST.get("file_id") or None
        include_subtree = request.POST.get("include_subtree") == "true"

        if user_id and role_id:
            try:
                if role_id == admin_role_id:
                    # Admin → assign everything
                    for folder in Folder.objects.filter(is_active=True):
                        assign_recursive(user_id, folder.id, role_id)
                    messages.success(request, "Admin role assigned with full access.")
                else:
                    if file_id and str(file_id).startswith("subfolder-"):
                        subfolder_id = int(file_id.replace("subfolder-", ""))
                        if include_subtree:
                            assign_recursive(user_id, subfolder_id, role_id)
                            messages.success(request, "Role assigned recursively to subfolder and its contents.")
                        else:
                            # Only assign to this subfolder
                            FolderUserRole.objects.get_or_create(
                                user_id=user_id,
                                folder_id=subfolder_id,
                                role_id=role_id,
                                defaults={"assigned_at": now()}
                            )
                            messages.success(request, "Role assigned to selected subfolder.")
                    else:
                        # Document or folder only
                        exists = FolderUserRole.objects.filter(
                            user_id=user_id,
                            folder_id=folder_id,
                            file_id=file_id,
                            role_id=role_id
                        ).exists()
                        if exists:
                            messages.warning(request, "This role assignment already exists.")
                        else:
                            FolderUserRole.objects.create(
                                user_id=user_id,
                                folder_id=folder_id,
                                file_id=file_id,
                                role_id=role_id,
                                assigned_at=now()
                            )
                            messages.success(request, "Role assigned successfully to document/folder.")

            except Exception as e:
                messages.error(request, f"Error: {str(e)}")
        else:
            messages.error(request, "User and role are required.")

        return redirect('assign_folder_role')

    # GET request → show form only
    users = User.objects.all()
    folders = Folder.objects.filter(is_active=True, parent__isnull=True)
    roles = Role.objects.all()
    documents = Document.objects.filter(is_deleted=False).select_related("folder")
    subfolders = Folder.objects.filter(is_active=True, parent__isnull=False)

    return render(request, "assign_folder_roles.html", {
        "users": users,
        "folders": folders,
        "roles": roles,
        "documents": documents,
        "subfolders": subfolders,
    })

def assign_recursive(user_id, folder_id, role_id):
    # Assign to the folder itself
    FolderUserRole.objects.get_or_create(user_id=user_id, folder_id=folder_id, role_id=role_id)

    # Assign to all documents in this folder
    documents = Document.objects.filter(folder_id=folder_id, is_deleted=False)
    for doc in documents:
        FolderUserRole.objects.get_or_create(user_id=user_id, file_id=doc.id, folder_id=folder_id, role_id=role_id)

    # Recurse into subfolders
    subfolders = Folder.objects.filter(parent_id=folder_id, is_active=True)
    for sub in subfolders:
        assign_recursive(user_id, sub.id, role_id)


from django.contrib.auth.decorators import login_required
from django.shortcuts import render

# @login_required
# def under_head_role(request):
#     # Show data only for GET request
#     if request.method == "GET":
#         # Get current user
#         current_user = request.user
#
#         # Fetch only those assignments where current user is the folder head
#         # Assuming Role model has a name field and folder head is stored as 'Folder Head' or similar
#         folder_head_roles = FolderUserRole.objects.filter(
#             user=current_user,
#             role__name__iexact="Folder Head"   # case-insensitive match
#         ).values_list('folder_id', flat=True)
#
#         # Fetch all assignments under those folders
#         assignments = FolderUserRole.objects.select_related("user", "folder", "role", "file").filter(
#             folder_id__in=folder_head_roles
#         )
#
#         # Other dropdown data
#         users = User.objects.exclude(is_superuser=True)
#         folders = Folder.objects.filter(id__in=folder_head_roles)
#         roles = Role.objects.all()
#         documents = Document.objects.select_related('folder').filter(folder_id__in=folder_head_roles)
#
#         return render(request, "assign_folder_roles.html", {
#             "users": users,
#             "folders": folders,
#             "roles": roles,
#             "assignments": assignments,
#             "documents": documents
#         })
#
#     # Optional: Handle POST or others if needed
#     else:
#         return render(request, "error.html", {"message": "Invalid request method."})

@login_required
def under_head_role(request):
    if request.method == "GET":
        # Get current user
        current_user = request.user

        folder_head_roles = FolderUserRole.objects.filter(
            user=current_user,
            role__role_name__iexact="Head"
        ).values_list('folder_id', flat=True)

        assignments = FolderUserRole.objects.select_related("user", "folder", "role", "file").filter(folder_id__in=folder_head_roles)
        return render(request, "folder_role_assignments.html", {"assignments": assignments})
    else:
        return render(request, "error.html", {"message": "Invalid request method."})


@login_required
def folder_role_assignments(request):
    assignments = FolderUserRole.objects.select_related("user", "folder", "role", "file").all()
    return render(request, "folder_role_assignments.html", {"assignments": assignments})

    # GET request
    users = User.objects.exclude(is_superuser=True)  # remove admin from user dropdown
    folders = Folder.objects.all()
    roles = Role.objects.all()
    documents = Document.objects.select_related('folder').all()
    assignments = FolderUserRole.objects.select_related('user', 'folder', 'role', 'file')

    return render(request, "assign_folder_roles.html", {
        "users": users,
        "folders": folders,
        "roles": roles,
        "assignments": assignments,
        "documents": documents
    })

@login_required
def remove_role(request, user_role_id):
    role_instance = get_object_or_404(FolderUserRole, id=user_role_id)

    # Check if removing last admin
    if role_instance.role.role_name.lower() == "admin":
        admin_count = FolderUserRole.objects.filter(role__role_name__iexact="admin").count()
        if admin_count <= 1:
            # First confirmation
            confirm1 = request.GET.get("confirm1")
            confirm2 = request.GET.get("confirm2")

            if not confirm1:
                messages.warning(request, "Are you sure you want to remove the LAST admin? Click again to confirm.")
                return redirect(f"{request.path}?confirm1=true")

            if not confirm2:
                messages.warning(request, "This is your FINAL confirmation. Click again to permanently remove.")
                return redirect(f"{request.path}?confirm1=true&confirm2=true")

    role_instance.delete()
    messages.success(request, "Role removed successfully.")
    return redirect("folder_role_assignments")

@login_required
def edit_list(request):
    user = request.user

    # Check if user is admin
    is_admin = FolderUserRole.objects.filter(
        user=user,
        role__role_name__iexact="admin"
    ).exists()

    if is_admin:
        # Admin sees all root folders
        folders = Folder.objects.filter(parent__isnull=True, is_active=True)
    else:
        # Non-admin sees only root folders they have access to
        folders = Folder.objects.filter(
            id__in=FolderUserRole.objects.filter(user=user).values_list('folder_id', flat=True),
            parent__isnull=True,
            is_active=True
        ).distinct()

    return render(request, 'edit_list.html', {'folders': folders})


@login_required
def folder_documents_edit(request, folder_id):
    # Get the current folder
    folder = get_object_or_404(Folder, id=folder_id, is_active=True)

    # Get documents directly under this folder
    documents = Document.objects.filter(
        folder=folder,
        is_deleted=False,
        folder__is_active=True
    )

    # Get subfolders directly under this folder
    subfolders = Folder.objects.filter(
        parent=folder,
        is_active=True
    )

    # Optional: Get all documents inside subfolders (non-recursive)
    subfolder_documents = Document.objects.filter(
        folder__in=subfolders,
        is_deleted=False,
        folder__is_active=True
    )

    return render(request, 'folder_documents.html', {
        'folder': folder,
        'documents': documents,
        'subfolders': subfolders,
        'subfolder_documents': subfolder_documents
    })


def edit_dynamic_table(request, doc_id):
    document = get_object_or_404(Document, id=doc_id, is_deleted=False)

    # Parse dynamic_data JSON
    dynamic_data = {}
    if document.dynamic_data:
        try:
            dynamic_data = json.loads(document.dynamic_data)
        except json.JSONDecodeError:
            dynamic_data = {}

    # The table name is the first key in the dict
    table_name = next(iter(dynamic_data.keys()), "")
    columns = dynamic_data.get(table_name, {}).get("columns", [])

    if request.method == "POST":
        # Handle remove column action
        if "remove_column" in request.POST:
            remove_index = int(request.POST.get("remove_column"))
            if 0 <= remove_index < len(columns):
                columns.pop(remove_index)

            # Save the updated JSON after deletion
            updated_data = {
                table_name: {
                    "columns": columns
                }
            }
            document.dynamic_data = json.dumps(updated_data, indent=4)
            document.save()

            return redirect("edit_dynamic_table", doc_id=doc_id)

        # Handle update action
        new_table_name = request.POST.get("table_name")
        new_columns = []
        total = int(request.POST.get("total_columns", 0))
        for i in range(total):
            col_name = request.POST.get(f"col_name_{i}")
            col_type = request.POST.get(f"col_type_{i}")
            col_constraints = request.POST.get(f"col_constraints_{i}", "")
            if col_name and col_type:
                col_data = {"name": col_name, "type": col_type}
                if col_constraints:
                    col_data["constraints"] = col_constraints
                new_columns.append(col_data)

        updated_data = {
            new_table_name: {
                "columns": new_columns
            }
        }

        document.dynamic_data = json.dumps(updated_data, indent=4)
        document.title = request.POST.get("title", document.title)
        document.description = request.POST.get("description", document.description)
        document.save()

        return redirect("edit_dynamic_table", doc_id=doc_id)

    return render(request, "edit_dynamic_table.html", {
        "document": document,
        "table_name": table_name,
        "columns": columns,
    })

@login_required
def history_index(request):
    user = request.user

    # Get user's role name
    try:
        role_name = UserRole.objects.select_related('role').get(user=user).role.role_name.lower()
    except UserRole.DoesNotExist:
        role_name = None

    if role_name == "admin":
        files = Upload.objects.select_related('uploaded_by', 'folder', 'document').filter(
            folder__is_active=True,
            document__is_deleted=False
        ).order_by('-upload_time')

    else:
        # Get folder IDs user has access to
        folder_ids = FolderUserRole.objects.filter(
            user=user,
            folder__isnull=False,
            folder__is_active=True
        ).values_list('folder_id', flat=True)

        # Get document IDs user has access to
        file_ids = FolderUserRole.objects.filter(
            user=user,
            file__isnull=False,
            file__is_deleted=False
        ).values_list('file_id', flat=True)

        # Query uploads where either:
        #   - Document is in allowed file_ids
        #   - Folder is in allowed folder_ids
        files = Upload.objects.select_related('uploaded_by', 'folder', 'document').filter(
            Q(document_id__in=file_ids) | Q(folder_id__in=folder_ids),
            folder__is_active=True,
            document__is_deleted=False
        ).order_by('-upload_time')

    return render(request, 'history_index.html', {'files': files})

@login_required
def history_view_file(request, file_id):
    file = get_object_or_404(
        Upload,
        id=file_id,
        folder__is_active=True,
        document__is_deleted=False
    )

    filename = file.file_name.lower()
    headers, rows = [], []
    preview_type, text_content, images_data = None, "", []

    try:
        if filename.endswith(".csv"):
            # ✅ Read CSV
            content = file.file_blob.decode("utf-8")
            reader = list(csv.reader(io.StringIO(content)))
            if reader:
                headers = reader[0]
                rows = reader[1:]

        elif filename.endswith(".xlsx"):
            # ✅ Read Excel
            in_memory_file = io.BytesIO(file.file_blob)
            wb = openpyxl.load_workbook(in_memory_file, data_only=True)
            sheet = wb.active
            data = list(sheet.iter_rows(values_only=True))

            if data:
                headers = [str(cell) if cell is not None else "" for cell in data[0]]
                rows = [
                    [str(cell) if cell is not None else "" for cell in row]
                    for row in data[1:]
                ]

        elif filename.endswith(".pdf"):
            # ✅ Mark for PDF preview
            preview_type = "pdf"

        elif filename.endswith(".docx"):
            # ✅ Extract text and images from DOCX
            in_memory_file = io.BytesIO(file.file_blob)
            doc = DocxDocument(in_memory_file)

            # Extract text
            text_content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

            # Extract images as base64
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_stream = rel.target_part.blob
                    b64 = base64.b64encode(image_stream).decode('utf-8')
                    mime = rel.target_part.content_type  # e.g., image/png
                    images_data.append(f"data:{mime};base64,{b64}")

            preview_type = "docx"

        elif filename.endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp")):
            # ✅ Image preview
            preview_type = "image"
        else:
            return render(request, "history_view.html", {
                "file": file,
                "headers": None,
                "rows": None,
                "preview_type": None,
                "message": "Preview not supported for this file type."
            })

        return render(request, "history_view.html", {
            "file": file,
            "headers": headers,
            "rows": rows,
            "preview_type": preview_type,
            "text_content": text_content,
            "images_data": images_data
        })

    except Exception as e:
        return HttpResponse(f"Error while reading file: {e}")


@login_required
def view_pdf(request, file_id):
    """Serve raw PDF so PDF.js can render it in canvas"""
    file = get_object_or_404(Upload, id=file_id)
    return FileResponse(io.BytesIO(file.file_blob), content_type="application/pdf")

@login_required
def history_edit_file(request, file_id):
    file_obj = get_object_or_404(
        Upload,
        id=file_id,
        folder__is_active=True,
        document__is_deleted=False
    )

    extension = file_obj.file_name.lower().split('.')[-1]

    try:
        file_data = file_obj.file_blob
        file_stream = BytesIO(file_data)

        if extension == 'csv':
            df = pd.read_csv(file_stream)
            if df.empty:
                return render(request, 'history_edit.html', {'error': 'File is empty or unreadable.'})
            context = {
                'df': df.to_dict(orient='records'),
                'columns': df.columns,
                'file_id': file_id,
                'file_type': 'csv'
            }

        elif extension in ['xls', 'xlsx']:
            df = pd.read_excel(file_stream, engine='openpyxl')
            if df.empty:
                return render(request, 'history_edit.html', {'error': 'File is empty or unreadable.'})
            context = {
                'df': df.to_dict(orient='records'),
                'columns': df.columns,
                'file_id': file_id,
                'file_type': 'excel'
            }

        else:
            return render(request, 'history_edit.html', {'error': 'Unsupported file format'})

        # ✅ Activity log
        ActivityLog.objects.create(
            user=request.user,
            upload=file_obj,
            file_name=file_obj.file_name,
            action="EDIT"
        )

        return render(request, 'history_edit.html', context)

    except Exception as e:
        return render(request, 'history_edit.html', {
            'error': f'❌ Error reading file: {e}'
        })

@login_required
def history_save_file(request, file_id):
    file = get_object_or_404(
        Upload,
        id=file_id,
        folder__is_active=True,
        document__is_deleted=False
    )

    extension = file.file_name.lower().split('.')[-1]

    try:
        if extension in ['csv', 'xls', 'xlsx']:
            # ✅ Recreate DataFrame from POST data
            headers = request.POST.getlist('columns')
            total_rows = int(request.POST.get('total_rows', 0))

            rows = []
            for i in range(total_rows):
                row = [request.POST.get(f'{col}_{i}', '') for col in headers]
                rows.append(row)

            df = pd.DataFrame(rows, columns=headers)

            # ✅ Save updated file_blob
            if extension == 'csv':
                file.file_blob = df.to_csv(index=False, header=True).encode('utf-8')
            else:  # Excel (xls / xlsx)
                output = BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, header=True)
                file.file_blob = output.getvalue()

            file.save()

            # ✅ Instead of going back → show updated preview
            return redirect('history_edit_file', file_id=file.id)

        else:
            return redirect('history_edit_file', file_id=file.id)

    except Exception as e:
        print("Error while saving file:", e)
        return redirect('history_edit_file', file_id=file.id)


def history_download_file(request, file_id):
    file_obj = get_object_or_404(Upload, id=file_id)

    # Ensure file has data
    if not file_obj.file_blob:
        raise Http404("File is missing.")

    content_type = file_obj.mime_type or 'application/octet-stream'

    response = HttpResponse(file_obj.file_blob, content_type=content_type)
    response['Content-Disposition'] = f'attachment; filename="{file_obj.file_name or "download"}"'
    return response


@login_required
def history_delete_file(request, file_id):
    file_obj = get_object_or_404(Upload, id=file_id)
    file_obj.is_deleted = True
    file_obj.save()
    ActivityLog.objects.create(
        user=request.user,
        upload=file_obj,
        file_name=file_obj.file_name,
        action="DELETE",
        timestamp=timezone.now()
    )

    return redirect('history_index')


@login_required
def edit_folders(request):
    if request.method == "POST":
        action = request.POST.get("action")

        if action == "add":
            folder_name = request.POST.get("folder_name")
            category = request.POST.get("category")
            if folder_name:
                Folder.objects.create(
                    folder_name=folder_name,
                    category=category or None,
                    parent=None,  # root folder
                    created_at=timezone.now(),
                    updated_at=timezone.now(),
                    is_active=True,
                )
                messages.success(request, "Folder added successfully.")
            else:
                messages.error(request, "Folder name cannot be empty.")


        elif action == "delete":

            folder_id = request.POST.get("folder_id")

            folder = Folder.objects.filter(id=folder_id, is_active=True).first()

            if folder:

                soft_delete_folder_and_descendants(folder)

                messages.success(request, "Folder, its subfolders and all related documents soft-deleted successfully.")

            else:
                messages.error(request, "Folder not found or already deleted.")

        return redirect("edit_folders")

    # Only show active root folders
    folders = Folder.objects.filter(is_active=True, parent__isnull=True).order_by("folder_name")
    return render(request, "edit_folders.html", {"folders": folders})

def collect_descendant_folder_ids(root_id):
    """
    Return list of descendant folder IDs (does NOT include root_id).
    Uses iterative BFS-style queries to avoid deep recursion and to be efficient.
    """
    descendants = []
    queue = [root_id]

    while True:
        children = list(Folder.objects.filter(parent_id__in=queue).values_list('id', flat=True))
        # stop when no more children
        if not children:
            break
        descendants.extend(children)
        queue = children

    return descendants  # list of ids (may be empty)

def soft_delete_folder_and_descendants(root_folder):
    """
    Soft-delete root_folder and every descendant folder + all documents under them.
    Uses bulk updates for efficiency. Does NOT call model .save() on each instance,
    so model signals (pre_save/post_save) will NOT be triggered.
    """
    now_ts = timezone.now()
    # collect descendant ids (does not include root)
    descendant_ids = collect_descendant_folder_ids(root_folder.id)
    # build set of all ids including root
    all_folder_ids = set(descendant_ids) | {root_folder.id}

    with transaction.atomic():
        # mark folders inactive and set updated_at
        Folder.objects.filter(id__in=all_folder_ids).update(is_active=False, updated_at=now_ts)

        # mark all docs in those folders as deleted
        Document.objects.filter(folder_id__in=all_folder_ids, is_deleted=False).update(is_deleted=True)


@login_required
def edit_folder_documents(request, folder_id):
    folder = get_object_or_404(Folder, id=folder_id, is_active=True)

    if request.method == "POST":
        action = request.POST.get("action")

        # ---------------- Add Document ----------------
        if action == "add_document":
            title = request.POST.get("title")
            description = request.POST.get("description", "")
            if title:
                Document.objects.create(
                    title=title,
                    description=description,
                    folder=folder,
                    dynamic_data="{}",  # empty JSON
                    created_at=timezone.now(),
                    is_deleted=False  # default to active
                )
                messages.success(request, "Document added successfully.")
            else:
                messages.error(request, "Document title cannot be empty.")

        # ---------------- Delete Document (Soft) ----------------
        elif action == "delete_document":
            doc_id = request.POST.get("doc_id")
            doc = Document.objects.filter(id=doc_id, folder=folder, is_deleted=False).first()
            if doc:
                doc.is_deleted = True
                doc.save(update_fields=["is_deleted"])
                messages.success(request, "Document marked as deleted.")
            else:
                messages.error(request, "Document not found or already deleted.")

        # ---------------- Add Subfolder ----------------
        elif action == "add_subfolder":
            subfolder_name = request.POST.get("subfolder_name")
            category = request.POST.get("category")
            if subfolder_name:
                Folder.objects.create(
                    folder_name=subfolder_name,
                    parent=folder,
                    category=category or None,
                    created_at=timezone.now(),
                    updated_at=timezone.now(),
                    is_active=True,
                )
                messages.success(request, "Subfolder added successfully.")
            else:
                messages.error(request, "Subfolder name cannot be empty.")

        # ---------------- Delete Subfolder (Recursive Soft Delete) ----------------
        elif action == "delete_subfolder":
            sub_id = request.POST.get("sub_id")
            # ensure subfolder is a child of current folder
            subfolder = Folder.objects.filter(id=sub_id, parent=folder).first()
            if subfolder and subfolder.is_active:
                soft_delete_folder_and_descendants(subfolder)
                messages.success(request, "Subfolder and all nested subfolders/documents soft-deleted successfully.")
            else:
                messages.error(request, "Subfolder not found or already deleted.")

    # ---------------- Display ----------------
    documents = Document.objects.filter(folder=folder, is_deleted=False)
    subfolders = Folder.objects.filter(parent=folder, is_active=True).order_by("folder_name")

    return render(request, "edit_folder_documents.html", {
        "folder": folder,
        "documents": documents,
        "subfolders": subfolders,
    })



from .models import ActivityLog

@login_required
def activity_log_view(request):

    logs = ActivityLog.objects.select_related("user", "upload").order_by("-timestamp")
    return render(request, "activity_log.html", {"logs": logs})

#
# @login_required
# def activity_log_view(request):
#     logs = ActivityLog.objects.select_related("user", "upload").order_by("-timestamp")
#
#     return render(request, "activity_log.html", {
#         "logs": logs
#     })

def history_restore_file(request, file_id):
    file_obj = get_object_or_404(Upload, id=file_id)
    file_obj.is_deleted = False
    file_obj.upload_time = timezone.now()   # 👈 restore hone par naya time de do
    file_obj.save()

    ActivityLog.objects.create(
        user=request.user,
        upload=file_obj,
        file_name=file_obj.file_name,
        action="RESTORE",
        timestamp=timezone.now()
    )
    return redirect('history_index')

def uploaded_files(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    files = Upload.objects.filter(document=document, is_deleted=False)

    return render(request, 'uploaded_files.html', {
        'document': document,
        'files': files
    })


# views.py

@login_required
def view_image(request, file_id):
    file = get_object_or_404(Upload, id=file_id)
    mime_type, _ = mimetypes.guess_type(file.file_name)
    return HttpResponse(file.file_blob, content_type=mime_type)

@login_required
def file_preview(request, file_id=None):
    folders = Folder.objects.filter(is_active=True).prefetch_related("upload_set")
    file = None
    headers, rows = [], []
    preview_type, text_content, images_data = None, "", []

    if file_id:
        file = get_object_or_404(
            Upload,
            id=file_id,
            folder__is_active=True,
            document__is_deleted=False
        )

        filename = file.file_name.lower()

        try:
            if filename.endswith(".csv"):
                content = file.file_blob.decode("utf-8")
                reader = list(csv.reader(io.StringIO(content)))
                if reader:
                    headers = reader[0]
                    rows = reader[1:]

            elif filename.endswith(".xlsx"):
                in_memory_file = io.BytesIO(file.file_blob)
                wb = openpyxl.load_workbook(in_memory_file, data_only=True)
                sheet = wb.active
                data = list(sheet.iter_rows(values_only=True))

                if data:
                    headers = [str(cell) if cell is not None else "" for cell in data[0]]
                    rows = [
                        [str(cell) if cell is not None else "" for cell in row]
                        for row in data[1:]
                    ]

            elif filename.endswith(".pdf"):
                preview_type = "pdf"

            elif filename.endswith(".docx"):
                in_memory_file = io.BytesIO(file.file_blob)
                doc = DocxDocument(in_memory_file)

                text_content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        image_stream = rel.target_part.blob
                        b64 = base64.b64encode(image_stream).decode('utf-8')
                        mime = rel.target_part.content_type
                        images_data.append(f"data:{mime};base64,{b64}")

                preview_type = "docx"

            elif filename.endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp")):
                preview_type = "image"

        except Exception as e:
            return HttpResponse(f"Error while reading file: {e}")

    return render(request, "preview.html", {
        "folders": folders,
        "file": file,
        "headers": headers,
        "rows": rows,
        "preview_type": preview_type,
        "text_content": text_content,
        "images_data": images_data,
    })
